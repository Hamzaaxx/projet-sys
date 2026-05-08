"""
Convert the corrected canaryfs report (markdown) into a properly formatted Word document.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent
MD_FILE = ROOT / "Compte_rendu_SE_corrige.md"
OUT_FILE = ROOT / "Compte_rendu_SE_corrige.docx"


def add_page_border(doc):
    """Add a thin page border to every page (cosmetic)."""
    for section in doc.sections:
        sectPr = section._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        for side in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '24')
            border.set(qn('w:color'), '888888')
            pgBorders.append(border)
        sectPr.append(pgBorders)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_inline(paragraph, text):
    """Render inline markdown: **bold**, *italic*, `code`."""
    pos = 0
    pattern = re.compile(r'(\*\*[^*]+?\*\*|\*[^*]+?\*|`[^`]+?`)')
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith('**') and token.endswith('**'):
            r = paragraph.add_run(token[2:-2]); r.bold = True
        elif token.startswith('*') and token.endswith('*'):
            r = paragraph.add_run(token[1:-1]); r.italic = True
        elif token.startswith('`') and token.endswith('`'):
            r = paragraph.add_run(token[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_table(lines, idx):
    """Parse a markdown table starting at idx. Returns (rows, lines_consumed)."""
    rows = []
    while idx < len(lines) and '|' in lines[idx]:
        line = lines[idx].strip()
        if line.startswith('|'): line = line[1:]
        if line.endswith('|'): line = line[:-1]
        cells = [c.strip() for c in line.split('|')]
        # skip the separator row (---|---|---)
        if all(re.match(r'^:?-+:?$', c) for c in cells):
            idx += 1
            continue
        rows.append(cells)
        idx += 1
    return rows, idx


def convert():
    doc = Document()

    # ── Default style ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ── Margins ──
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Footer with page number ──
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run().add_text("ENSET Mohammedia 2025-2026 — canaryfs — page ")
    # Page number field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run = fp.add_run()
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    text = MD_FILE.read_text(encoding='utf-8')
    lines = text.split('\n')

    i = 0
    in_code = False
    code_buf = []
    code_lang = ''

    while i < len(lines):
        line = lines[i]

        # ── code blocks ──
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_lang = line.strip()[3:]
                code_buf = []
            else:
                # close code block
                in_code = False
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                # background shading
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), 'F4F4F4')
                pPr.append(shd)
                code_text = '\n'.join(code_buf)
                r = p.add_run(code_text)
                r.font.name = 'Consolas'
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        stripped = line.strip()

        # ── horizontal rule = page break ──
        if stripped == '---':
            doc.add_page_break()
            i += 1
            continue

        # ── headings ──
        if stripped.startswith('# '):
            p = doc.add_heading(stripped[2:].strip(), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            i += 1; continue
        if stripped.startswith('## '):
            t = stripped[3:].strip()
            p = doc.add_heading(t, level=1)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
            i += 1; continue
        if stripped.startswith('### '):
            p = doc.add_heading(stripped[4:].strip(), level=2)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
            i += 1; continue

        # ── tables ──
        if '|' in line and i + 1 < len(lines) and re.match(r'^\s*\|?[\s\-:|]+\|?\s*$', lines[i + 1]):
            rows, new_i = parse_table(lines, i)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Light Grid Accent 1'
                for r, row in enumerate(rows):
                    for c, cell_txt in enumerate(row):
                        cell = table.cell(r, c)
                        cell.text = ''
                        p = cell.paragraphs[0]
                        add_inline(p, cell_txt)
                        if r == 0:
                            for run in p.runs:
                                run.bold = True
                            shade_cell(cell, '2E74B5')
                            for run in p.runs:
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                doc.add_paragraph()
            i = new_i
            continue

        # ── bullet list ──
        if re.match(r'^\s*[-*]\s+', line):
            content = re.sub(r'^\s*[-*]\s+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, content)
            i += 1; continue

        # ── numbered list ──
        if re.match(r'^\s*\d+\.\s+', line):
            content = re.sub(r'^\s*\d+\.\s+', '', line)
            p = doc.add_paragraph(style='List Number')
            add_inline(p, content)
            i += 1; continue

        # ── empty line ──
        if stripped == '':
            i += 1; continue

        # ── normal paragraph ──
        p = doc.add_paragraph()
        add_inline(p, line)
        i += 1

    add_page_border(doc)
    doc.save(OUT_FILE)
    print(f"Saved: {OUT_FILE}")


if __name__ == '__main__':
    convert()
