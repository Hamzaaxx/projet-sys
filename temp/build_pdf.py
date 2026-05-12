"""
Build a Word document matching the original ENSET PDF structure,
then convert it to PDF using Word.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx2pdf import convert

ROOT = Path(__file__).parent
MD_FILE = ROOT / "Compte_rendu_SE_corrige.md"
DOCX_FILE = ROOT / "Compte_rendu_SE_v2.docx"
PDF_FILE = ROOT / "Compte_rendu_SE_v2.pdf"


# ──────────────────────────────────────────────────────────────
# helpers
# ──────────────────────────────────────────────────────────────
def set_cell_border(cell, color="1F4E79", size="12"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_inline(paragraph, text):
    pos = 0
    pattern = re.compile(r'(\*\*[^*]+?\*\*|`[^`]+?`)')
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith('**'):
            r = paragraph.add_run(token[2:-2]); r.bold = True
        elif token.startswith('`'):
            r = paragraph.add_run(token[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_table(lines, idx):
    rows = []
    while idx < len(lines) and '|' in lines[idx]:
        line = lines[idx].strip()
        if line.startswith('|'): line = line[1:]
        if line.endswith('|'): line = line[:-1]
        cells = [c.strip() for c in line.split('|')]
        if all(re.match(r'^:?-+:?$', c) for c in cells):
            idx += 1
            continue
        rows.append(cells)
        idx += 1
    return rows, idx


# ──────────────────────────────────────────────────────────────
# Cover page — matches the original ENSET layout
# ──────────────────────────────────────────────────────────────
def build_cover_page(doc):
    # Top header table — 3 columns: French left, logo center, Arabic right
    header = doc.add_table(rows=1, cols=3)
    header.autofit = False
    widths = [Cm(6.5), Cm(3.5), Cm(6.5)]
    for i, w in enumerate(widths):
        header.columns[i].width = w
    header.cell(0, 0).width = widths[0]
    header.cell(0, 1).width = widths[1]
    header.cell(0, 2).width = widths[2]

    # Left: French institutional name
    left = header.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = left.add_run("Ecole Normale Supérieure de\nl'Enseignement Technique\nMohammedia")
    r.bold = True; r.font.size = Pt(9)
    p2 = header.cell(0, 0).add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r2 = p2.add_run("Université Hassan II de Casablanca")
    r2.bold = True; r2.font.size = Pt(8)

    # Center: ENSET text logo
    center = header.cell(0, 1).paragraphs[0]
    center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = center.add_run("\nE N S E T")
    rc.bold = True; rc.font.size = Pt(20); rc.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # Right: Arabic
    right = header.cell(0, 2).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = right.add_run("المدرسة العليا لأساتذة التعليم التقني\nالمحمدية")
    rr.bold = True; rr.font.size = Pt(10)
    p3 = header.cell(0, 2).add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr2 = p3.add_run("جامعة الحسن الثاني بالدار البيضاء")
    rr2.bold = True; rr2.font.size = Pt(9)

    # Spacer
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(20)

    # DEPARTEMENT line
    dept = doc.add_paragraph()
    dept.alignment = WD_ALIGN_PARAGRAPH.LEFT
    dr = dept.add_run("DEPARTEMENT MATHEMATIQUES ET INFORMATIQUE")
    dr.bold = True; dr.font.size = Pt(13); dr.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # Spacer
    doc.add_paragraph().paragraph_format.space_before = Pt(20)

    # Title block
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = p1.add_run("Compte rendu du")
    rt.bold = True; rt.font.size = Pt(28); rt.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    p2t = doc.add_paragraph()
    p2t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt2 = p2t.add_run("Projet Pratique")
    rt2.bold = True; rt2.font.size = Pt(28); rt2.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    doc.add_paragraph().paragraph_format.space_before = Pt(15)

    # Filière box (blue border)
    filiere_table = doc.add_table(rows=1, cols=1)
    fcell = filiere_table.cell(0, 0)
    set_cell_border(fcell, color="1F4E79", size="12")
    fp = fcell.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = fp.add_run("Filière :\nGLSID-ICCN")
    rf.bold = True; rf.font.size = Pt(18); rf.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    fcell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph().paragraph_format.space_before = Pt(15)

    # Project title — double blue border
    title_table = doc.add_table(rows=1, cols=1)
    tcell = title_table.cell(0, 0)
    set_cell_border(tcell, color="1F4E79", size="18")
    tp = tcell.paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt3 = tp.add_run(
        "CanaryFS : Système de surveillance par fichiers leurres "
        "pour la détection et la traçabilité des accès non autorisés"
    )
    rt3.bold = True; rt3.font.size = Pt(14); rt3.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph().paragraph_format.space_before = Pt(20)

    # Authors / Supervisor — two-column table
    rs_table = doc.add_table(rows=1, cols=2)
    rs_table.autofit = False
    rs_table.columns[0].width = Cm(8.5)
    rs_table.columns[1].width = Cm(8.0)
    rs_table.cell(0, 0).width = Cm(8.5)
    rs_table.cell(0, 1).width = Cm(8.0)

    set_cell_border(rs_table.cell(0, 0), color="1F4E79", size="8")
    set_cell_border(rs_table.cell(0, 1), color="1F4E79", size="8")

    # Left cell — authors
    left_cell = rs_table.cell(0, 0)
    p_realise = left_cell.paragraphs[0]
    rrl = p_realise.add_run("Réalisé par :")
    rrl.bold = True; rrl.underline = True; rrl.font.size = Pt(11)
    for name in ["Douae TAHIRI", "Hamza BELAZRI", "Youssef SABRI", "Mohamed AIT EL KADI"]:
        np = left_cell.add_paragraph(name)
        np.runs[0].font.size = Pt(11)

    # Right cell — supervisor
    right_cell = rs_table.cell(0, 1)
    p_enc = right_cell.paragraphs[0]
    rer = p_enc.add_run("Encadré par :")
    rer.bold = True; rer.underline = True; rer.font.size = Pt(11)
    np2 = right_cell.add_paragraph("M. OUAGUID")
    np2.runs[0].font.size = Pt(11)

    doc.add_paragraph().paragraph_format.space_before = Pt(20)

    # Year
    year_p = doc.add_paragraph()
    year_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ryr = year_p.add_run("Année Universitaire : 2025 - 2026")
    ryr.bold = True; ryr.font.size = Pt(13)

    doc.add_paragraph().paragraph_format.space_before = Pt(60)

    # Footer block (school address)
    footer_t = doc.add_table(rows=1, cols=1)
    fct = footer_t.cell(0, 0)
    shade_cell(fct, "1F4E79")
    fp = fct.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf1 = fp.add_run("ENSET, Avenue Hassan II - B.P. 159 - Mohammedia - Maroc")
    rf1.bold = True; rf1.font.size = Pt(9); rf1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    fp2 = fct.add_paragraph()
    fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf2 = fp2.add_run("☎ 05 23 32 22 20 / 05 23 32 35 30 – Fax : 05 23 32 25 46 - Site Web: www.enset-media.ac.ma")
    rf2.font.size = Pt(8); rf2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    fp3 = fct.add_paragraph()
    fp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf3 = fp3.add_run("E-Mail : contact@enset-media.ac.ma")
    rf3.font.size = Pt(8); rf3.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_page_break()


# ──────────────────────────────────────────────────────────────
# Body builder — converts markdown body into Word
# ──────────────────────────────────────────────────────────────
def build_body(doc, md_text):
    lines = md_text.split('\n')
    # find where the cover content ends — skip everything before first ---
    started = False
    skip_until_first_hr = True
    i = 0
    in_code = False
    code_buf = []

    # Skip the cover-page markdown (we built it manually)
    # The markdown has: # Compte rendu... up to first --- = cover page
    # Then Remerciements section starts
    while i < len(lines):
        if lines[i].strip() == '---':
            i += 1
            break
        i += 1

    while i < len(lines):
        line = lines[i]

        # code blocks
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), 'F5F5F5')
                pPr.append(shd)
                r = p.add_run('\n'.join(code_buf))
                r.font.name = 'Consolas'
                r.font.size = Pt(9)
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        s = line.strip()

        if s == '---':
            doc.add_page_break()
            i += 1; continue

        # H1 = Chapitre — page break before
        if s.startswith('# '):
            t = s[2:].strip()
            doc.add_page_break()
            p = doc.add_heading('', level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(t)
            r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            i += 1; continue

        # H2
        if s.startswith('## '):
            t = s[3:].strip()
            p = doc.add_heading('', level=2)
            r = p.add_run(t)
            r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            i += 1; continue

        # H3
        if s.startswith('### '):
            t = s[4:].strip()
            p = doc.add_heading('', level=3)
            r = p.add_run(t)
            r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            i += 1; continue

        # tables
        if '|' in line and i + 1 < len(lines) and re.match(r'^\s*\|?[\s\-:|]+\|?\s*$', lines[i + 1]):
            rows, new_i = parse_table(lines, i)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Light Grid Accent 1'
                for r_idx, row in enumerate(rows):
                    for c, cell_txt in enumerate(row):
                        cell = table.cell(r_idx, c)
                        cell.text = ''
                        p = cell.paragraphs[0]
                        add_inline(p, cell_txt)
                        if r_idx == 0:
                            for run in p.runs:
                                run.bold = True
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            shade_cell(cell, '1F4E79')
                doc.add_paragraph()
            i = new_i
            continue

        # bullets
        if re.match(r'^\s*[-*]\s+', line):
            content = re.sub(r'^\s*[-*]\s+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, content)
            i += 1; continue

        # numbered
        if re.match(r'^\s*\d+\.\s+', line):
            content = re.sub(r'^\s*\d+\.\s+', '', line)
            p = doc.add_paragraph(style='List Number')
            add_inline(p, content)
            i += 1; continue

        if s == '':
            i += 1; continue

        # paragraph
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(8)
        add_inline(p, line)
        i += 1


# ──────────────────────────────────────────────────────────────
# Build the document
# ──────────────────────────────────────────────────────────────
def main():
    doc = Document()

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Footer with page numbers (skip on cover page)
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run = fp.add_run()
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    # Cover page
    build_cover_page(doc)

    # Body content
    md_text = MD_FILE.read_text(encoding='utf-8')
    build_body(doc, md_text)

    # Save .docx
    doc.save(DOCX_FILE)
    print(f"Saved DOCX: {DOCX_FILE}")

    # Convert to PDF
    print("Converting to PDF (Word will open briefly)...")
    convert(str(DOCX_FILE), str(PDF_FILE))
    print(f"Saved PDF:  {PDF_FILE}")


if __name__ == '__main__':
    main()
